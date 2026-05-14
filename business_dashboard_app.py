import streamlit as st
import pandas as pd
import plotly.express as px

from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import letter
from io import BytesIO
from docx import Document
from docx.shared import Inches
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    Image
)
st.set_page_config(
    page_title="E&E ProfitMatrix",
    page_icon="📊",
    layout="wide"
)

try:
    APP_PASSWORD = st.secrets["APP_PASSWORD"]
except Exception:
    APP_PASSWORD = "ProfitPilot2026"

password = st.text_input("Enter password", type="password")

if password != APP_PASSWORD:
    st.warning("Please enter the correct password to access ProfitMatrix.")
    st.stop()

st.markdown(
    """
    <style>

    /* MAIN BACKGROUND */
    .stApp {
        background:
        radial-gradient(circle at top left, #1e3a8a15, transparent 25%),
        radial-gradient(circle at bottom right, #2563eb10, transparent 30%),
        linear-gradient(180deg, #f8fafc 0%, #eef2ff 100%);
    }

    /* SIDEBAR */
    section[data-testid="stSidebar"] {
        background: linear-gradient(180deg, #0f172a 0%, #1e293b 100%);
        border-right: 1px solid #334155;
    }

    section[data-testid="stSidebar"] * {
        color: white !important;
    }

    /* KPI CARDS */
    div[data-testid="stMetric"] {
        background: rgba(255,255,255,0.92);
        backdrop-filter: blur(12px);
        padding: 24px;
        border-radius: 22px;
        box-shadow:
            0 10px 30px rgba(15, 23, 42, 0.10),
            0 2px 6px rgba(15, 23, 42, 0.05);
        border: 1px solid rgba(255,255,255,0.4);
        transition: all 0.25s ease;
    }

    div[data-testid="stMetric"]:hover {
        transform: translateY(-4px);
        box-shadow:
            0 14px 34px rgba(15, 23, 42, 0.15),
            0 4px 10px rgba(15, 23, 42, 0.08);
    }

    div[data-testid="stMetricLabel"] {
        font-size: 15px;
        color: #475569;
        font-weight: 600;
    }

    div[data-testid="stMetricValue"] {
        font-size: 38px;
        font-weight: 800;
        color: #0f172a;
    }

    /* TABS */
    .stTabs [data-baseweb="tab-list"] {
        gap: 14px;
        border-bottom: none;
    }

    .stTabs [data-baseweb="tab"] {
        background: rgba(255,255,255,0.75);
        backdrop-filter: blur(10px);
        border-radius: 16px 16px 0 0;
        padding: 12px 22px;
        font-weight: 700;
        border: 1px solid rgba(255,255,255,0.35);
        transition: all 0.25s ease;
    }

    .stTabs [aria-selected="true"] {
        background: linear-gradient(135deg,#2563eb,#1d4ed8);
        color: white !important;
        box-shadow: 0 8px 24px rgba(37,99,235,0.25);
    }

    /* FILE UPLOADER */
    section[data-testid="stFileUploader"] {
        background: rgba(255,255,255,0.82);
        backdrop-filter: blur(12px);
        border-radius: 24px;
        padding: 18px;
        border: 1px solid rgba(255,255,255,0.45);
        box-shadow:
            0 8px 26px rgba(15, 23, 42, 0.08);
    }

    /* BUTTONS */
    .stButton button,
    .stDownloadButton button {
        border-radius: 14px;
        padding: 12px 22px;
        font-weight: 700;
        border: none;
        background: linear-gradient(135deg,#2563eb,#1d4ed8);
        color: white;
        box-shadow: 0 6px 18px rgba(37,99,235,0.25);
        transition: all 0.25s ease;
    }

    .stButton button:hover,
    .stDownloadButton button:hover {
        transform: translateY(-2px);
        box-shadow: 0 10px 24px rgba(37,99,235,0.35);
    }

    /* ALERT BOXES */
    .stAlert {
        border-radius: 18px;
        border: none;
        box-shadow: 0 6px 20px rgba(15,23,42,0.06);
    }

    /* TABLES */
    .stDataFrame {
        border-radius: 18px;
        overflow: hidden;
        box-shadow: 0 8px 22px rgba(15,23,42,0.08);
    }

    /* HEADERS */
    h1, h2, h3 {
        color: #0f172a;
        font-weight: 800;
    }

    /* FOOTER */
    footer {
        visibility: hidden;
    }

    </style>
    """,
    unsafe_allow_html=True
)

st.markdown(
    """
    <div style="display:flex; align-items:center; gap:18px;">
        <div style="
            width:70px;
            height:70px;
            border-radius:18px;
            background:linear-gradient(135deg,#1f2937,#2563eb);
            display:flex;
            align-items:center;
            justify-content:center;
            color:white;
            font-size:24px;
            font-weight:800;">
            E&E
        </div>
        <div>
            <h1>E&E ProfitMatrix Business Dashboard</h1>
            <p style="font-size:18px; margin-top:6px;">
                Business intelligence for sales, profit, margins and product risk analysis.
            </p>
        </div>
    </div>
    """,
    unsafe_allow_html=True
)
def create_word_report():
    buffer = BytesIO()
    doc = Document()

    doc.add_heading("ProfitPilot Executive Business Report", 0)

    doc.add_paragraph(
        "This report provides an automated business intelligence analysis based on the uploaded sales data. "
        "It evaluates revenue, costs, profit, margins, product-level risks, and strategic opportunities."
    )
    doc.add_heading('Sales Performance Chart', level=2)
    doc.add_picture("sales_chart.png", width=Inches(6))

    doc.add_heading('Product Profitability Chart', level=2)
    doc.add_picture("product_chart.png", width=Inches(6))
    doc.add_heading("1. Executive Summary", level=1)

    doc.add_paragraph(
        f"The business generated total revenue of ${total_revenue:,.2f}, "
        f"with total costs of ${total_cost:,.2f}, resulting in net profit of ${total_profit:,.2f}. "
        f"The average profit margin was {avg_margin:.2f}%."
    )

    if avg_margin >= 40:
        doc.add_paragraph(
            "Overall profitability appears strong. The business is generating a healthy margin relative to its cost structure."
        )
    elif avg_margin >= 15:
        doc.add_paragraph(
            "Overall profitability is positive but should be monitored. Some pricing or cost adjustments may improve performance."
        )
    else:
        doc.add_paragraph(
            "Profitability is weak. The business should urgently review pricing, costs, and product mix."
        )

    doc.add_heading("2. Key Performance Indicators", level=1)

    doc.add_paragraph(f"Total Revenue: ${total_revenue:,.2f}")
    doc.add_paragraph(f"Total Cost: ${total_cost:,.2f}")
    doc.add_paragraph(f"Total Profit: ${total_profit:,.2f}")
    doc.add_paragraph(f"Average Margin: {avg_margin:.2f}%")

    doc.add_heading("3. Product Profitability Analysis", level=1)

    best_product = product_analysis.sort_values("profit", ascending=False).iloc[0]
    worst_product = product_analysis.sort_values("profit").iloc[0]

    doc.add_paragraph(
        f"The best-performing product was {best_product['product']}, "
        f"with profit of ${best_product['profit']:,.2f}. "
        "This product should be prioritized in sales, promotion, and inventory planning."
    )

    doc.add_paragraph(
        f"The weakest product was {worst_product['product']}, "
        f"with profit of ${worst_product['profit']:,.2f}. "
        "This product should be reviewed for pricing, supplier cost, demand, or operational inefficiency."
    )

    doc.add_heading("4. Risk Alerts", level=1)

    loss_products = product_analysis[product_analysis["profit"] < 0]
    low_margin_products = product_analysis[
        (product_analysis["margin"] < 15) & (product_analysis["profit"] >= 0)
    ]

    if len(loss_products) > 0:
        doc.add_paragraph(
            f"The analysis identified {len(loss_products)} product(s) generating losses. "
            "These products reduce total profitability and require immediate attention."
        )
    else:
        doc.add_paragraph("No products with negative profit were detected.")

    if len(low_margin_products) > 0:
        doc.add_paragraph(
            f"The analysis identified {len(low_margin_products)} product(s) with low margins. "
            "These products may still generate sales but contribute weakly to profitability."
        )
    else:
        doc.add_paragraph("No critical low-margin products were detected.")

    doc.add_heading("5. Product Risk Table", level=1)

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"

    headers = table.rows[0].cells
    headers[0].text = "Product"
    headers[1].text = "Revenue"
    headers[2].text = "Cost"
    headers[3].text = "Profit"
    headers[4].text = "Margin"
    headers[5].text = "Status"

    for _, row in product_analysis.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["product"])
        cells[1].text = f"${row['revenue']:,.2f}"
        cells[2].text = f"${row['cost']:,.2f}"
        cells[3].text = f"${row['profit']:,.2f}"
        cells[4].text = f"{row['margin']:.2f}%"
        cells[5].text = str(row["status"])

    doc.add_heading("6. Strategic Recommendations", level=1)

    doc.add_paragraph(
        "1. Prioritize products with strong margins and positive profit contribution."
    )
    doc.add_paragraph(
        "2. Review products classified as Loss and identify whether the problem comes from price, cost, demand, or waste."
    )
    doc.add_paragraph(
        "3. Evaluate low-margin products and consider price adjustments, supplier negotiation, or product replacement."
    )
    doc.add_paragraph(
        "4. Use this dashboard regularly to monitor business performance instead of relying only on total sales."
    )

    doc.add_heading("7. Business Interpretation", level=1)

    doc.add_paragraph(
        "A business can show positive sales while still losing money on specific products. "
        "ProfitPilot helps identify where profit is actually being generated and where hidden losses may exist. "
        "This makes the tool useful for pricing decisions, inventory control, cost reduction, and product strategy."
    )

    doc.save(buffer)
    buffer.seek(0)
    return buffer
uploaded_file = st.file_uploader(
    "📤 Upload your Excel or CSV file",
    type=["xlsx", "csv"]
)


def load_file(uploaded_file):
 from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    Image
)

from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import letter

def create_pdf_report():

    buffer = BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=40,
        leftMargin=40,
        topMargin=40,
        bottomMargin=40
    )

    styles = getSampleStyleSheet()
    elements = []

    title = Paragraph(
        "<b>ProfitPilot Executive Business Report</b>",
        styles["Title"]
    )
    elements.append(title)
    elements.append(Spacer(1, 16))

    intro = Paragraph(
        "This report provides an automated business intelligence analysis based on the uploaded sales data. "
        "It evaluates revenue, costs, profit, margins, product-level risks, and strategic opportunities.",
        styles["BodyText"]
    )
    elements.append(intro)
    elements.append(Spacer(1, 18))

    elements.append(Paragraph("<b>1. Executive Summary</b>", styles["Heading1"]))

    summary = Paragraph(
        f"The business generated total revenue of <b>${total_revenue:,.2f}</b>, "
        f"with total costs of <b>${total_cost:,.2f}</b>, resulting in net profit of "
        f"<b>${total_profit:,.2f}</b>. The average profit margin was "
        f"<b>{avg_margin:.2f}%</b>.",
        styles["BodyText"]
    )
    elements.append(summary)
    elements.append(Spacer(1, 12))

    if avg_margin >= 40:
        margin_text = "Overall profitability appears strong. The business is generating a healthy margin relative to its cost structure."
    elif avg_margin >= 15:
        margin_text = "Overall profitability is positive but should be monitored. Some pricing or cost adjustments may improve performance."
    else:
        margin_text = "Profitability is weak. The business should urgently review pricing, costs, and product mix."

    elements.append(Paragraph(margin_text, styles["BodyText"]))
    elements.append(Spacer(1, 18))

    elements.append(Paragraph("<b>2. Key Performance Indicators</b>", styles["Heading1"]))

    kpi_data = [
        ["Metric", "Value"],
        ["Total Revenue", f"${total_revenue:,.2f}"],
        ["Total Cost", f"${total_cost:,.2f}"],
        ["Total Profit", f"${total_profit:,.2f}"],
        ["Average Margin", f"{avg_margin:.2f}%"]
    ]

    kpi_table = Table(kpi_data, colWidths=[230, 230])
    kpi_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2563eb")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("BACKGROUND", (0, 1), (-1, -1), colors.whitesmoke),
        ("PADDING", (0, 0), (-1, -1), 8)
    ]))

    elements.append(kpi_table)
    elements.append(Spacer(1, 20))

    elements.append(Paragraph("<b>3. Sales Performance Chart</b>", styles["Heading1"]))
    elements.append(Image("sales_chart.png", width=460, height=260))
    elements.append(Spacer(1, 20))

    elements.append(Paragraph("<b>4. Product Profitability Chart</b>", styles["Heading1"]))
    elements.append(Image("product_chart.png", width=460, height=260))
    elements.append(Spacer(1, 20))

    best_product = product_analysis.sort_values("profit", ascending=False).iloc[0]
    worst_product = product_analysis.sort_values("profit").iloc[0]

    elements.append(Paragraph("<b>5. Product Profitability Analysis</b>", styles["Heading1"]))

    elements.append(Paragraph(
        f"The best-performing product was <b>{best_product['product']}</b>, "
        f"with profit of <b>${best_product['profit']:,.2f}</b>. "
        "This product should be prioritized in sales, promotion, and inventory planning.",
        styles["BodyText"]
    ))

    elements.append(Spacer(1, 10))

    elements.append(Paragraph(
        f"The weakest product was <b>{worst_product['product']}</b>, "
        f"with profit of <b>${worst_product['profit']:,.2f}</b>. "
        "This product should be reviewed for pricing, supplier cost, demand, or operational inefficiency.",
        styles["BodyText"]
    ))

    elements.append(Spacer(1, 18))

    elements.append(Paragraph("<b>6. Risk Alerts</b>", styles["Heading1"]))

    loss_products = product_analysis[product_analysis["profit"] < 0]
    low_margin_products = product_analysis[
        (product_analysis["margin"] < 15) & (product_analysis["profit"] >= 0)
    ]

    if len(loss_products) > 0:
        elements.append(Paragraph(
            f"The analysis identified <b>{len(loss_products)}</b> product(s) generating losses. "
            "These products reduce total profitability and require immediate attention.",
            styles["BodyText"]
        ))
    else:
        elements.append(Paragraph(
            "No products with negative profit were detected.",
            styles["BodyText"]
        ))

    elements.append(Spacer(1, 8))

    if len(low_margin_products) > 0:
        elements.append(Paragraph(
            f"The analysis identified <b>{len(low_margin_products)}</b> product(s) with low margins. "
            "These products may still generate sales but contribute weakly to profitability.",
            styles["BodyText"]
        ))
    else:
        elements.append(Paragraph(
            "No critical low-margin products were detected.",
            styles["BodyText"]
        ))

    elements.append(Spacer(1, 18))

    elements.append(Paragraph("<b>7. Product Risk Table</b>", styles["Heading1"]))

    table_data = [["Product", "Revenue", "Cost", "Profit", "Margin", "Status"]]

    for _, row in product_analysis.iterrows():
        table_data.append([
            str(row["product"]),
            f"${row['revenue']:,.2f}",
            f"${row['cost']:,.2f}",
            f"${row['profit']:,.2f}",
            f"{row['margin']:.2f}%",
            str(row["status"])
        ])

    risk_table = Table(table_data, colWidths=[95, 75, 75, 75, 70, 95])
    risk_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f2937")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("BACKGROUND", (0, 1), (-1, -1), colors.whitesmoke),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("PADDING", (0, 0), (-1, -1), 5)
    ]))

    elements.append(risk_table)
    elements.append(Spacer(1, 20))

    elements.append(Paragraph("<b>8. Strategic Recommendations</b>", styles["Heading1"]))

    recommendations = [
        "Prioritize products with strong margins and positive profit contribution.",
        "Review products classified as Loss and identify whether the problem comes from price, cost, demand, or waste.",
        "Evaluate low-margin products and consider price adjustments, supplier negotiation, or product replacement.",
        "Use this dashboard regularly to monitor business performance instead of relying only on total sales."
    ]

    for rec in recommendations:
        elements.append(Paragraph(f"• {rec}", styles["BodyText"]))
        elements.append(Spacer(1, 6))

    elements.append(Spacer(1, 12))

    elements.append(Paragraph("<b>9. Business Interpretation</b>", styles["Heading1"]))

    elements.append(Paragraph(
        "A business can show positive sales while still losing money on specific products. "
        "ProfitPilot helps identify where profit is actually being generated and where hidden losses may exist. "
        "This makes the tool useful for pricing decisions, inventory control, cost reduction, and product strategy.",
        styles["BodyText"]
    ))

    doc.build(elements)

    buffer.seek(0)
    return buffer


def load_file(uploaded_file):
    if uploaded_file.name.endswith(".csv"):
        return pd.read_csv(uploaded_file)

    raw_excel = pd.read_excel(uploaded_file, header=None)

    keywords = [
        "fecha", "date",
        "producto", "product",
        "cantidad", "quantity",
        "precio", "price",
        "costo", "cost"
    ]

    header_row = None

    for i in range(len(raw_excel)):
        row_values = raw_excel.iloc[i].fillna("").astype(str).str.lower().tolist()
        row_text = " ".join(row_values)

        matches = sum(keyword in row_text for keyword in keywords)

        if matches >= 2:
            header_row = i
            break

    if header_row is None:
        st.error("Could not detect the header row in this Excel file.")
        st.write("Detected preview:")
        st.dataframe(raw_excel.head(15), use_container_width=True)
        st.stop()

    return pd.read_excel(uploaded_file, header=header_row)


if uploaded_file is not None:

    df = load_file(uploaded_file)

    df.columns = (
        df.columns
        .astype(str)
        .str.strip()
        .str.lower()
        .str.replace(" ", "_")
    )

    COLUMN_MAPPINGS = {

    "date": [
        "date", "fecha", "day"
    ],

    "product": [
        "product", "producto", "service",
        "servicio", "item", "category"
    ],

    "quantity": [
        "quantity", "cantidad", "qty",
        "units"
    ],

    "unit_price": [
        "unit_price", "price", "precio",
        "sale_price", "revenue"
    ],

    "unit_cost": [
        "unit_cost", "cost", "costo",
        "expense"
    ],

    "employee": [
        "employee", "empleado",
        "worker", "staff",
        "barber", "lawyer"
    ]
}
    def detect_column(df_columns, possible_names):

     for col in df_columns:

        normalized = col.lower().strip()

        if normalized in possible_names:
            return col

     return None
    
    column_mapping = {
        "fecha": "date",
        "date": "date",

        "producto": "product",
        "product": "product",
        "item": "product",
        "articulo": "product",
        "artículo": "product",

        "cantidad": "quantity",
        "quantity": "quantity",
        "qty": "quantity",
        "unidades": "quantity",
        "units": "quantity",

        "precio": "unit_price",
        "precio_unitario": "unit_price",
        "price": "unit_price",
        "unit_price": "unit_price",

        "costo": "unit_cost",
        "costo_unitario": "unit_cost",
        "cost": "unit_cost",
        "unit_cost": "unit_cost"
    }

    df.rename(columns=column_mapping, inplace=True)

    required_columns = ["date", "product", "quantity", "unit_price", "unit_cost"]
    missing_columns = [col for col in required_columns if col not in df.columns]

    if missing_columns:
        st.error(f"Missing required columns: {missing_columns}")
        st.write("Detected columns:")
        st.write(list(df.columns))
        st.dataframe(df.head(10), use_container_width=True)
        st.stop()

    df["date"] = pd.to_datetime(df["date"], errors="coerce")
    df["quantity"] = pd.to_numeric(df["quantity"], errors="coerce")
    df["unit_price"] = pd.to_numeric(df["unit_price"], errors="coerce")
    df["unit_cost"] = pd.to_numeric(df["unit_cost"], errors="coerce")

    df = df.dropna(subset=["date", "product", "quantity", "unit_price", "unit_cost"])
    date_col = detect_column(
    df.columns,
    COLUMN_MAPPINGS["date"]
)

    product_col = detect_column(
    df.columns,
    COLUMN_MAPPINGS["product"]
)

    quantity_col = detect_column(
    df.columns,
    COLUMN_MAPPINGS["quantity"]
)

    price_col = detect_column(
    df.columns,
    COLUMN_MAPPINGS["unit_price"]
)

    cost_col = detect_column(
    df.columns,
    COLUMN_MAPPINGS["unit_cost"]
)

    employee_col = detect_column(
    df.columns,
    COLUMN_MAPPINGS["employee"]
)

    df["revenue"] = (
    df[quantity_col] * df[price_col]
)

    df["cost"] = (
    df[quantity_col] * df[cost_col]
)

    df["profit"] = (
    df["revenue"] - df["cost"]
)

    df["margin"] = (
    df["profit"] / df["revenue"]
) * 100

    total_revenue = df["revenue"].sum()
    total_cost = df["cost"].sum()
    total_profit = df["profit"].sum()
    avg_margin = df["margin"].mean()

    product_analysis = df.groupby(product_col, as_index=False).agg({
    "revenue": "sum",
    "cost": "sum",
    "profit": "sum"
})

    product_analysis["margin"] = (
        product_analysis["profit"] / product_analysis["revenue"]
    ) * 100

    def classify_product(profit, margin):
        if profit < 0:
            return "🔴 Loss"
        elif margin < 15:
            return "🟡 Low Margin"
        else:
            return "🟢 Healthy"

    product_analysis["status"] = product_analysis.apply(
        lambda row: classify_product(row["profit"], row["margin"]),
        axis=1
    )
    product_text = " ".join(
        product_analysis[product_col].astype(str).str.lower().tolist()
    )

    business_type = "General Business"

    if any(word in product_text for word in [
        "coffee", "burger", "pizza", "sandwich", "juice", "bread", "cake", "restaurant", "food"
    ]):
        business_type = "Food / Restaurant Business"

    elif any(word in product_text for word in [
        "hair", "barber", "beard", "manicure", "pedicure", "nails", "tint", "tinte", "corte", "barba"
    ]):
        business_type = "Beauty / Barber Business"

    elif any(word in product_text for word in [
        "case", "legal", "lawyer", "contract", "consulting", "consultoria", "abogado", "caso"
    ]):
        business_type = "Legal / Professional Services"

    elif any(word in product_text for word in [
        "project", "construction", "cement", "labor", "material", "obra", "proyecto", "construccion"
    ]):
        business_type = "Construction / Project Business"

    elif any(word in product_text for word in [
        "membership", "gym", "training", "fitness", "coach", "membresia"
    ]):
        business_type = "Fitness / Membership Business"

    with st.sidebar:

     st.markdown("## 🚀 E&E ProfitMatrix")

    st.markdown("""
    Business Intelligence Platform
    """)

    st.divider()

    st.markdown("### 📌 Features")

    st.markdown("""
    - Sales analytics
    - Margin analysis
    - Product risk detection
    - Executive reports
    - CSV exports
    - PDF & Word reports
    """)

    st.divider()

    st.markdown("### 📈 Platform Status")

    st.success("System operational")

    st.divider()

    st.markdown("""
    **Developed by E&E Analytics**
    
    Business Intelligence & Data Solutions
    """
    )
    tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "📌 Executive Summary",
    "📈 Charts",
    "⚠️ Product Risk",
    "📥 Downloads",
    "🤖 AI Insights"
])

    with tab1:
        col1, col2, col3, col4 = st.columns(4)

        col1.metric("Revenue", f"${total_revenue:,.2f}")
        col2.metric("Cost", f"${total_cost:,.2f}")
        col3.metric("Profit", f"${total_profit:,.2f}")
        col4.metric("Margin", f"{avg_margin:.2f}%")
        st.info(f"🏢 Detected business type: **{business_type}**")

        st.subheader("Business Alerts")

        if (product_analysis["profit"] < 0).any():
            st.error("⚠️ Some products are generating losses.")

        if (product_analysis["margin"] < 15).any():
            st.warning("⚠️ Some products have low margins.")

        if not (product_analysis["profit"] < 0).any() and not (product_analysis["margin"] < 15).any():
            st.success("✅ Business performance looks healthy.")

    with tab2:
        daily = df.groupby("date", as_index=False)[["revenue", "cost", "profit"]].sum()

        fig = px.line(
            daily,
            x="date",
            y=["revenue", "cost", "profit"],
            markers=True,
            title="Revenue, Cost and Profit Over Time",
            color_discrete_sequence=["#2563eb", "#dc2626", "#16a34a"]
        )
        # fig.write_image("sales_chart.png")
        st.plotly_chart(fig, use_container_width=True)

        fig2 = px.bar(
            product_analysis,
            x="product",
            y="profit",
            color="status",
            text_auto=True,
            title="Profit by Product",
            color_discrete_map={
                "🟢 Healthy": "#16a34a",
                "🟡 Low Margin": "#eab308",
                "🔴 Loss": "#dc2626"
                }
        )
        # fig2.write_image("product_chart.png")
        st.plotly_chart(fig2, use_container_width=True)

    with tab3:
        st.subheader("Product Risk Analysis")
        st.dataframe(product_analysis, use_container_width=True)

        st.subheader("Detailed Data")
        st.dataframe(df, use_container_width=True)

    with tab4:
        analyzed_csv = df.to_csv(index=False).encode("utf-8")
        product_csv = product_analysis.to_csv(index=False).encode("utf-8")

        st.download_button(
            "📥 Download analyzed data CSV",
            data=analyzed_csv,
            file_name="profitpilot_analyzed_data.csv",
            mime="text/csv"
        )

        st.download_button(
            "📥 Download product risk analysis CSV",
            data=product_csv,
            file_name="profitpilot_product_risk.csv",
            mime="text/csv"
        )
        
        st.download_button(
            "📄 Download executive report Word",
            data=create_word_report(),
            file_name="profitpilot_executive_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        st.download_button(
            "📕 Download executive report PDF",
            data=create_pdf_report(),
            file_name="profitpilot_executive_report.pdf",
            mime="application/pdf"
)
    with tab5:
        st.subheader("🤖 Smart Business Insights")
        st.info(f"🏢 Detected business profile: **{business_type}**")
        
        top_revenue_item = product_analysis.sort_values(
            "revenue",
            ascending=False
        ).iloc[0]

        top_profit_item = product_analysis.sort_values(
            "profit",
            ascending=False
        ).iloc[0]

        lowest_margin_item = product_analysis.sort_values(
            "margin",
            ascending=True
        ).iloc[0]

        loss_count = (product_analysis["profit"] < 0).sum()
        low_margin_count = (product_analysis["margin"] < 15).sum()

        st.markdown("### 🔎 Key Findings")

        st.success(
            f"💰 Highest revenue item: **{top_revenue_item[product_col]}** "
            f"with ${top_revenue_item['revenue']:,.2f} in revenue."
        )

        st.success(
            f"🏆 Most profitable item: **{top_profit_item[product_col]}** "
            f"with ${top_profit_item['profit']:,.2f} in profit."
        )

        st.warning(
            f"⚠️ Lowest margin item: **{lowest_margin_item[product_col]}** "
            f"with {lowest_margin_item['margin']:.2f}% margin."
        )

        if loss_count > 0:
            st.error(
                f"🚨 {loss_count} item(s) are generating losses. "
                "These should be reviewed immediately."
            )
        else:
            st.success("✅ No loss-generating items were detected.")

        if low_margin_count > 0:
            st.warning(
                f"⚠️ {low_margin_count} item(s) have low margins. "
                "Consider price adjustments or cost reduction."
            )
        else:
            st.success("✅ Margin structure looks healthy overall.")

        st.markdown("### 🧠 Strategic Recommendation")

        if avg_margin >= 50:
            st.info(
                "The business shows strong profitability. The main strategy should be "
                "to scale high-margin products/services and maintain cost discipline."
            )
        elif avg_margin >= 25:
            st.info(
                "The business is profitable but has room for improvement. Focus on "
                "optimizing costs, reviewing prices, and promoting the most profitable items."
            )
        else:
            st.info(
                "The business margin is weak. Priority should be given to cost control, "
                "pricing review, and elimination or redesign of low-margin items."
            )
else:

    st.info("Upload your file to start.")

sample = pd.DataFrame({
        "Fecha": ["2026-01-01", "2026-01-02", "2026-01-03"],
        "Producto": ["Coffee", "Sandwich", "Juice"],
        "Cantidad": [30, 15, 20],
        "Precio Unitario": [4.5, 8.0, 5.0],
        "Costo Unitario": [1.2, 3.5, 2.0]
    })

st.subheader("Example format")
st.dataframe(sample, use_container_width=True)

st.markdown("---")

st.markdown(
    """
    <div style='text-align:center; color:gray; padding:20px;'>
        © 2026 E&E ProfitMatrix • Business Intelligence Platform
    </div>
    """,
    unsafe_allow_html=True
)